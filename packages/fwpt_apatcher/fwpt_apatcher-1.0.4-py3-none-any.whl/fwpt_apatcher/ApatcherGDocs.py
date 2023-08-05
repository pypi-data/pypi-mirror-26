import glob
import logging
import operator
import os
import re
import sys
from enum import Enum

from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Cm
from docxtpl import DocxTemplate

import fwpt_apatcher.ApatcherUtils as autil

logger = logging.getLogger(__name__)

PATH_TO_IMG = "cfg/box_img.jpg"
PATH_TO_TMPL = "cfg/_temp.sql"
PATH_TO_CFG = "cfg/config.ini"
WARNING_MSG_DELIMITER = ">>>"


class ExtNameDocTpl(Enum):
    CHANGELIST = "changelist"
    UPDATE_LOG = "updatelog"
    BUILDS_LIST = "buildlist"


class ProjectPatchEntity:
    version_num = None
    changed_files = []
    db_changed = []
    ui_changed = []
    comment = None
    author = None
    path_patch_ = None
    name_file_ = None

    def __init__(self, version_num=None, author=None, changed_files=None, ui_changed=None, db_changed=None,
                 comment=None):
        self.version_num = version_num
        self.author = author
        self.changed_files = changed_files
        self.ui_changed = ui_changed
        self.db_changed = db_changed
        self.comment = comment

    def initFromSource(self, path_patch_, name_file_):
        self.path_patch_ = path_patch_
        self.name_file_ = name_file_
        self.init_from_exists()

    def init_from_exists(self):
        full_txt = autil.get_patch_top_txt(self.path_patch_)

        remap = {ord('\t'): None, ord('\f'): None, ord('\r'): None}
        num_patch = 0

        try:
            author = full_txt[full_txt.find("Автор:") + len("Автор"): full_txt.find("Дата:")].strip(" :\n").translate(
                remap)
            num_patch = full_txt[
                        full_txt.find("Номер патча:") + len("Номер патча"): full_txt.find("Номер тикета:")].strip(
                " :\n").translate(remap)
            descr = full_txt[full_txt.find("Комментарий:") + len("Комментарий"): full_txt.find("Создан:")].strip(
                " :\n").replace("\n", "").translate(remap)
            if full_txt.find("Создан:") != -1 and full_txt.find("Список включённых файлов:") != -1:
                list_files = full_txt[full_txt.find("Список включённых файлов:") + len("Список включённых файлов"): len(
                    full_txt)].strip(" :\n").translate(remap)
                # удалим возможно предупреждение в шапке
                p_delim_place = list_files.find(WARNING_MSG_DELIMITER)
                if p_delim_place != -1:
                    list_files = list_files[:p_delim_place]
                lst_files = [x.strip(" ") for x in (list_files.lstrip(":")).split(",")]
            else:
                new_files = full_txt[full_txt.find("Новые объекты:") + len("Новые объекты"): full_txt.find(
                    "Измененные объекты:")].strip(" :\n").translate(remap)
                change_files = full_txt[full_txt.find("Измененные объекты:") + len("Измененные объекты"): full_txt.find(
                    "Удаленные объекты:")].strip(" :\n").translate(remap)
                del_files = full_txt[full_txt.find("Удаленные объекты:") + len("Удаленные объекты"): full_txt.find(
                    "Комментарий:")].strip(" :\n").translate(remap)
                list_files = new_files + change_files + del_files
                lst_files = [x.strip(" ") for x in (list_files.lstrip(":")).split(",")]

            self.version_num = num_patch.lstrip("0")
            self.changed_files = lst_files
            db_change, web_change = autil.split_list_files(self.changed_files)
            self.db_changed = db_change
            self.ui_changed = web_change
            self.comment = descr
            self.author = author
        except Exception as e:
            logging.error("Некорректный парсинг шапки патча. Исключение: {}".format(e))
            raise Exception("Некорректный парсинг шапки патча #{}".format(num_patch))

    def __str__(self):
        return "PatchEntity:\n\tversion={ver}\n\tcomment={comment}\n\tchanged={files}\n".format(ver=self.version_num,
                                                                                                comment=self.comment,
                                                                                                files=self.changed_files)


class ProjectCorpus:
    project_name = None
    project_ext = None
    patches = []
    weight = None
    num_patches = []
    path_to = None

    def __init__(self, project_name=None, project_ext=None, patches=None, path_to=None):
        self.project_name = project_name
        self.project_ext = project_ext
        if patches is not None and len(patches) > 0:
            for x in patches:
                if not isinstance(x, ProjectPatchEntity):
                    raise TypeError("An invalid type for the project patch collection.")
            self.patches = patches
        else:
            self.patches = []
        self.path_to = path_to

    def set_primary(self, project_name, project_ext, weight, num_patches):
        self.weight = weight
        self.num_patches = num_patches
        self.project_ext = project_ext
        self.project_name = project_name

    def get_objects_for_transfer(self):
        return [i.path_patch_ for i in self.patches]

    def prepare(self):
        path_patches = os.path.join(self.path_to, "patches")
        for p in self.num_patches:
            xname = path_patches + "\\**\\*_*{0}.sql".format(str(p))
            res_finding = glob.glob(xname, recursive=True)
            if res_finding is not None and len(res_finding) > 0:
                lst_tm = res_finding[0]
            else:
                logging.warning("Патч по маске={mask} будет пропущен".format(mask=xname))
                continue
            _ppe = ProjectPatchEntity()
            _ppe.initFromSource(path_patch_=lst_tm, name_file_=os.path.basename(lst_tm))
            self.patches.append(_ppe)

    def print_patches_info(self):
        for i in self.patches:
            print(i)

    def __str__(self):
        return "Project Corpus:\n\tproject={project}\n\tnum_patches={num}\n\tweight={w}\n\text={ext}\n\tpath={path}".format(
            project=self.project_name,
            num=self.num_patches,
            w=self.weight, ext=self.project_ext, path=self.path_to)


class DocEntity:
    doc_name = None
    ext = None
    date_str = None
    customer_dir = None
    print_author = None
    projects = []
    path_to = None

    def __init__(self, doc_name, ext, projects, customer_dir, print_author, date_str):
        self.doc_name = doc_name
        if not isinstance(ext, ExtNameDocTpl):
            raise TypeError("An invalid type for the ext attr of Doc entity object.")
        self.ext = ext
        for x in projects:
            if not isinstance(x, ProjectCorpus):
                raise TypeError("An invalid type for the project corpus in document.")
        self.projects = projects
        self.customer_dir = customer_dir
        self.print_author = print_author
        self.date_str = date_str

    def get_docx_tpl_name(self):
        if self.ext == ExtNameDocTpl.CHANGELIST:
            return "changelist_tpl.docx"
        elif self.ext == ExtNameDocTpl.UPDATE_LOG:
            return "update_log_tpl.docx"
        elif self.ext == ExtNameDocTpl.BUILDS_LIST:
            return "builds_changelist_tpl.docx"
        else:
            raise NotImplementedError("The requested document type is not implemented or found!")

    def generate_update_log_context(self, docx_tpl):
        document = docx_tpl.new_subdoc()
        table = document.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "№п.п"
        hdr_cells[1].text = "Наименование обновления (патча)"
        hdr_cells[2].text = "Результат установки (Ок/Err)"
        hdr_cells[3].text = "Сотрудник, проводивший установку"
        hdr_cells[4].text = "Комментарий"
        hdr_cells[0].width = Cm(1)
        hdr_cells[1].width = Cm(10)
        hdr_cells[2].width = Cm(1)
        hdr_cells[3].width = Cm(1)
        hdr_cells[4].width = Cm(7)
        list_patch = [jpatch.name_file_ for iproject in self.projects for jpatch in iproject.patches]

        for idx, item in enumerate(list_patch, start=1):
            row_cells = table.add_row().cells
            row_cells[0].text = str(idx)
            row_cells[1].text = item
            row_cells[2].text = ""
            row_cells[3].text = ""
            row_cells[4].text = ""

        context = {
            "date_p": self.date_str,
            "author_p": self.print_author,
            "path_customer_dir": self.customer_dir,
            "tab_updates": document
        }
        return context

    def generate_changelist_context(self, docx_tpl):
        document = docx_tpl.new_subdoc()

        for proj_idx, xproj in enumerate(self.projects, start=1):
            p_1 = document.add_paragraph()
            r_1 = p_1.add_run()
            r_1.add_break()
            p_1.add_run("Таблица №{numtab}. Описание обновлений ({pr_name} патчи)".format(numtab=proj_idx,
                                                                                          pr_name=xproj.project_name),
                        style='TitleStyle2').bold = True
            table = document.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "№ патча"
            hdr_cells[1].text = "Краткое описание обновления (патча)"
            hdr_cells[2].text = "Перечень изменяемых объектов в БД"
            hdr_cells[3].text = "Перечень изменяемых объектов в веб-интерфейсе"
            hdr_cells[0].width = Cm(1)
            hdr_cells[1].width = Cm(10)
            hdr_cells[2].width = Cm(10)
            hdr_cells[3].width = Cm(7)

            for ypatch in xproj.patches:
                row_cells = table.add_row().cells
                row_cells[0].text = ypatch.version_num
                row_cells[1].text = ypatch.comment
                row_cells[2].text = ",\n".join(map(str, ypatch.db_changed))
                row_cells[3].text = ",\n".join(map(str, ypatch.ui_changed))

        context = {
            "project_changes": document
        }
        return context

    def generate(self):
        tpl_name = os.path.join(os.path.dirname(os.path.realpath(sys.argv[0])), "cfg", self.get_docx_tpl_name())
        if not os.path.isfile(tpl_name):
            logging.error("{} error: Path {} not exists".format(self, tpl_name))
        docx_tpl = DocxTemplate(tpl_name)
        context = {}
        if self.ext == ExtNameDocTpl.UPDATE_LOG:
            context = self.generate_update_log_context(docx_tpl)
        elif self.ext == ExtNameDocTpl.CHANGELIST:
            context = self.generate_changelist_context(docx_tpl)
        docx_tpl.render(context)
        self.path_to = "tmp\\{}".format(self.doc_name)
        docx_tpl.save(self.path_to)

    def get_doc_path(self):
        return self.path_to

    def __str__(self):
        return "DocEntity: name={name}, ext={ext}".format(name=self.doc_name, ext=self.ext)


def get_project_ext_weight(project_ext):
    if project_ext == "sdk":
        s_w = 0
    elif project_ext == "base":
        s_w = 5
    elif project_ext.startswith("option_"):
        s_w = 10
    else:
        s_w = 1000
    return s_w


def parse_namespace_pc(s, root_path):
    l_pcs = []
    pre_pat_project_corpus_ = re.compile("([a-zA-Z_-]*?):([\d-]*?)[,\]]")
    for i in pre_pat_project_corpus_.findall(s):
        proj_ext_ = i[0]
        proj_name_ = proj_ext_.replace("option_", "").replace("_", " ").upper() if proj_ext_ != "base" else "Базовые"

        if proj_ext_ == "base":
            path_to = root_path
        elif proj_ext_ == "builds":
            path_to = ""
        elif proj_ext_.startswith("option_") or proj_ext_ == "sdk":
            path_to = os.path.join(root_path, proj_ext_)
        else:
            path_to = os.path.join(root_path, "projects", proj_ext_)

        _pc = ProjectCorpus(path_to=path_to)
        spp_ = i[1]
        try:
            pp = [int(spp_)] if len(spp_.split("-")) < 2 else list(
                range(int(spp_.split("-")[0]), int(spp_.split("-")[1]) + 1))
        except ValueError:
            pp = []
        _pc.set_primary(project_name=proj_name_, project_ext=proj_ext_, weight=get_project_ext_weight(proj_ext_),
                        num_patches=pp)
        if len(_pc.num_patches) > 0: l_pcs.append(_pc)
        l_pcs.sort(key=operator.attrgetter("weight"))
    return l_pcs
